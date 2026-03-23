"""
Document Parser for AI Customer Service System
Handles parsing of various document formats (PDF, DOCX, TXT)
"""

import logging
from typing import List, Optional
from pathlib import Path

# Import required libraries
try:
    import PyPDF2
    import pdfplumber
    from docx import Document
except ImportError as e:
    logging.error(f"Missing required packages for document parsing: {e}")
    raise


class DocumentParser:
    """
    Class for parsing different document formats
    """

    def __init__(self):
        """Initialize the document parser"""
        logging.info("DocumentParser initialized")

    def parse_pdf(self, file_path: str) -> str:
        """
        Parse PDF file and extract text content
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content from the PDF
        """
        try:
            text_content = ""
            
            # Using pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            logging.info(f"Successfully parsed PDF: {file_path}, extracted {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logging.error(f"Error parsing PDF {file_path}: {str(e)}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text_content = ""
                    for page in reader.pages:
                        text_content += page.extract_text() + "\n"
                
                logging.info(f"Fallback parsing successful for PDF: {file_path}")
                return text_content
            except Exception as fallback_error:
                logging.error(f"Fallback PDF parsing also failed: {fallback_error}")
                raise e

    def parse_docx(self, file_path: str) -> str:
        """
        Parse DOCX file and extract text content
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content from the DOCX
        """
        try:
            doc = Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            text_content = "\n".join(paragraphs)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += "\n" + cell.text
            
            logging.info(f"Successfully parsed DOCX: {file_path}, extracted {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logging.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise

    def parse_txt(self, file_path: str) -> str:
        """
        Parse TXT file and extract text content
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text content from the TXT
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            logging.info(f"Successfully parsed TXT: {file_path}, extracted {len(text_content)} characters")
            return text_content
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    text_content = file.read()
                
                logging.info(f"Successfully parsed TXT (GBK encoding): {file_path}, extracted {len(text_content)} characters")
                return text_content
            except Exception as e:
                logging.error(f"Error parsing TXT {file_path} with multiple encodings: {str(e)}")
                raise
        except Exception as e:
            logging.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise

    def parse_document(self, file_path: str) -> str:
        """
        Universal method to parse any supported document format
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.parse_pdf(file_path)
        elif file_extension == '.docx':
            return self.parse_docx(file_path)
        elif file_extension == '.txt':
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}. Supported types: .pdf, .docx, .txt")


# Initialize logging
logging.basicConfig(level=logging.INFO)


def test_document_parser():
    """
    Test function for the document parser
    """
    print("Testing Document Parser...")
    
    parser = DocumentParser()
    
    # Example usage would go here
    print("DocumentParser initialized successfully")


if __name__ == "__main__":
    test_document_parser()